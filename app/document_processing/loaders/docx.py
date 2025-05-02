"""
DOCX document loader.

This module provides a document loader for Microsoft Word DOCX files.
"""
import os
from pathlib import Path
from typing import List, Optional, Dict, Any
import uuid

import docx
from app.core.interfaces import RagDocument, DocumentType
from app.document_processing.loaders.base import BaseDocumentLoader, document_loader_registry
from app.utils.logging import log


class DocxDocumentLoader(BaseDocumentLoader):
    """Document loader for DOCX files."""
    
    def __init__(self, extract_headers: bool = True, extract_tables: bool = True):
        """
        Initialize the DOCX document loader.
        
        Args:
            extract_headers: Whether to extract headers and include them in the text.
            extract_tables: Whether to extract tables. Not yet fully implemented.
        """
        super().__init__()
        self.supported_extensions = {"docx"}
        self.extract_headers = extract_headers
        self.extract_tables = extract_tables
    
    def _load_file(self, file_path: str) -> List[RagDocument]:
        """
        Load a DOCX file.
        
        Args:
            file_path: Path to the DOCX file.
            
        Returns:
            List containing a single RagDocument with the document content.
        """
        path = Path(file_path)
        
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        if not path.is_file():
            raise ValueError(f"Not a file: {file_path}")
        
        try:
            log.debug(f"Loading DOCX file: {file_path}")
            
            # Load the document
            doc = docx.Document(file_path)
            
            # Extract document properties
            doc_properties = {}
            try:
                core_properties = doc.core_properties
                if core_properties:
                    for prop_name in dir(core_properties):
                        if not prop_name.startswith('_') and not callable(getattr(core_properties, prop_name)):
                            value = getattr(core_properties, prop_name)
                            if value is not None:
                                doc_properties[f"docx_{prop_name}"] = str(value)
            except Exception as e:
                log.warning(f"Error extracting document properties: {e}")
            
            # Extract text
            paragraphs = []
            
            # Extract headers if requested
            if self.extract_headers:
                for section in doc.sections:
                    header = section.header
                    if header:
                        for paragraph in header.paragraphs:
                            if paragraph.text.strip():
                                paragraphs.append(f"[HEADER] {paragraph.text}")
            
            # Extract main content paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    style_name = paragraph.style.name if paragraph.style else "Normal"
                    
                    # Handle headings differently
                    if style_name.startswith('Heading'):
                        level = style_name.replace('Heading', '').strip() or '1'
                        paragraphs.append(f"{'#' * int(level)} {paragraph.text}")
                    else:
                        paragraphs.append(paragraph.text)
            
            # Extract tables if requested (simplified approach)
            if self.extract_tables:
                for i, table in enumerate(doc.tables):
                    paragraphs.append(f"\n[TABLE {i+1}]\n")
                    for row in table.rows:
                        row_text = " | ".join(cell.text for cell in row.cells)
                        if row_text.strip():
                            paragraphs.append(row_text)
            
            # Combine everything into a single text
            content = "\n\n".join(paragraphs)
            
            # Create metadata
            metadata = self._create_metadata(
                source_path=file_path,
                **doc_properties
            )
            
            # Create document
            doc = RagDocument(
                content=content,
                metadata=metadata,
                doc_id=str(uuid.uuid4()),
                doc_type=DocumentType.DOCX
            )
            
            return [doc]
        
        except Exception as e:
            log.error(f"Error loading DOCX file {file_path}: {e}")
            raise


# Register the loader
document_loader_registry.register(DocxDocumentLoader()) 